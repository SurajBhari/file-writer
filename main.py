import wikipediaapi

from json import load, dump


from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE

schema = load(open('schema.json'))

wiki_wiki = wikipediaapi.Wikipedia('en')


document = Document()


heading_schema = schema["headings"]

wiki_link = str(input("Please input Topic EXACT wikipedia name\n"))
page_py = wiki_wiki.page(wiki_link)
print(page_py.fullurl)
if not page_py.exists():
    quit("Page Doesn't exist")

data = {"data":[]}

def get_format(sections, level=0):
    lis = []
    for s in sections:
        if len(s.text) < 500 and s.level == 0:
            continue # Take care of short passages
        print("%s: %s - %s" % ("*" * (level + 1), s.title, s.text[0:40]))
        x = {"title":s.title, "content":s.text, "level":level}
        x["section"] = get_format(s.sections, level + 1)
        lis.append(x)
    
    return lis

index = 0
data = get_format(page_py.sections)

with open("data.json", "w") as f:
    dump(data, f, indent=4)

document.add_heading(page_py.title, 0).style.font.size = Pt(schema["headings"]["fontsize"]*2)
for topic in data:
    #Make Page Head
    document.add_heading(topic["title"]).style.font.size = Pt(schema["headings"]["fontsize"])
    
    
    font = document.add_paragraph(topic["content"]).style.font
    font.size = Pt(schema["paragraph"]["fontsize"])
    font.bold = schema["paragraph"]["bold"]
    font.italic = schema["paragraph"]["italic"]
    font.underline = schema["paragraph"]["underline"]
    font.strike = schema["paragraph"]["strikethrough"]
    
    for section in topic["section"]:
        h = document.add_paragraph(section["title"]) #Disguise it as a paragraph so its easier to modify
        h.style.font.size = Pt(schema["subtopics"]["heading_fontsize"])
        h_format = h.style.paragraph_format
        h_format.left_indent = Inches(0.25)
        h_format.space_before = Pt(12)
        h_format.widow_control = True
        
        p = document.add_paragraph(section["content"])
        p_format = p.style.paragraph_format
        p_format.left_indent = Inches(0.25)
        p_format.space_before = Pt(12)
        p_format.widow_control = True

        p.style.font.size = Pt(schema["subtopics"]["fontsize"])
        p.style.font.bold = schema["subtopics"]["bold"]
        p.style.font.italic = schema["subtopics"]["italic"]
        p.style.font.underline = schema["subtopics"]["underline"]
        p.style.font.strike = schema["subtopics"]["strikethrough"]
        
    document.add_page_break()
document.save('random_paragraph.docx')
